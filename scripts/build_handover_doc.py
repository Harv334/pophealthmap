#!/usr/bin/env python3
"""Write the handover Word document.

    python scripts/build_handover_doc.py "PopHealth Map - Handover.docx"

The same ground as HANDOVER.md, in the format people actually forward to each
other and read on a train. It is generated rather than hand-written so it can
be produced again when the repository moves on, and so the tables of what
refreshes and what does not stay in one place.

Needs python-docx:  pip install python-docx
"""
from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Inches

NHS_BLUE = RGBColor(0x00, 0x5E, 0xB8)
NHS_DARK = RGBColor(0x00, 0x30, 0x87)
INK = RGBColor(0x21, 0x2B, 0x32)
GREY = RGBColor(0x4C, 0x62, 0x72)


def main() -> int:
    out = Path(sys.argv[1] if len(sys.argv) > 1 else "PopHealth Map - Handover.docx")
    d = Document()

    base = d.styles["Normal"]
    base.font.name = "Calibri"
    base.font.size = Pt(11)
    base.font.color.rgb = INK
    base.paragraph_format.space_after = Pt(8)

    def h(text, level=1):
        p = d.add_heading(text, level=level)
        for r in p.runs:
            r.font.color.rgb = NHS_DARK if level == 1 else NHS_BLUE
        return p

    def para(text, bold=False, italic=False, size=11, colour=INK):
        p = d.add_paragraph()
        r = p.add_run(text)
        r.bold, r.italic = bold, italic
        r.font.size = Pt(size)
        r.font.color.rgb = colour
        return p

    def bullet(text):
        return d.add_paragraph(text, style="List Bullet")

    def numbered(text):
        return d.add_paragraph(text, style="List Number")

    def code(text):
        p = d.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(text)
        r.font.name = "Consolas"
        r.font.size = Pt(9.5)
        r.font.color.rgb = GREY
        return p

    def table(headers, rows, widths=None):
        t = d.add_table(rows=1, cols=len(headers))
        t.style = "Light Grid Accent 1"
        for i, htxt in enumerate(headers):
            cell = t.rows[0].cells[i]
            cell.text = ""
            r = cell.paragraphs[0].add_run(htxt)
            r.bold = True
            r.font.size = Pt(9.5)
        for row in rows:
            cells = t.add_row().cells
            for i, val in enumerate(row):
                cells[i].text = ""
                r = cells[i].paragraphs[0].add_run(str(val))
                r.font.size = Pt(9.5)
        if widths:
            for row in t.rows:
                for i, w in enumerate(widths):
                    row.cells[i].width = Inches(w)
        d.add_paragraph()
        return t

    # ── Title ────────────────────────────────────────────────────────────────
    t = d.add_heading("PopHealth Map", level=0)
    for r in t.runs:
        r.font.color.rgb = NHS_DARK
    sub = para("Handover notes: running it, refreshing it, and adding to it",
               italic=True, size=12, colour=GREY)
    sub.alignment = WD_ALIGN_PARAGRAPH.LEFT

    para("A map of population health across all 33 London boroughs. It is a "
         "static website plus one Python script that rebuilds it from published "
         "data. There is no server, no database and no login. GitHub hosts the "
         "site, GitHub Actions refreshes the data, and everything a visitor "
         "sees is a file in the repository.")

    # ── 1 ────────────────────────────────────────────────────────────────────
    h("1. Does anyone need to do anything?", 1)
    para("No. The data refreshes itself and the site publishes itself.", bold=True)
    para("On the 15th of each month GitHub Actions downloads every source, "
         "rebuilds the files the map reads, and commits them. That commit "
         "publishes the site. Nobody has to be there for it.")
    para("Three things are worth doing anyway:")
    table(["When", "What"],
          [("Monthly, one minute",
            "Open the Actions tab and check the last Refresh data run is green."),
           ("If a run fails",
            "It opens an issue naming the source that broke. data/meta/manifest.json "
            "records what happened to every source on the last run."),
           ("Once a year",
            "Read section 4, on what goes stale. Those are dates and identifiers "
            "that were true when written and will stop being true.")],
          widths=[1.7, 4.6])

    # ── 2 ────────────────────────────────────────────────────────────────────
    h("2. Running it", 1)

    h("Look at the site on your own machine", 2)
    para("The site is static, but it fetches data files, so opening index.html "
         "straight from the file system will not work. Serve it instead:")
    code("python -m http.server 8000")
    para("Then open http://127.0.0.1:8000/index.html")

    h("Publish it", 2)
    para("Every push to the main branch publishes. There is nothing to build "
         "and nothing to upload. The live site is whatever is on main.")

    h("Rebuild the data yourself", 2)
    code("pip install -r requirements.txt\npython fetch_all_data.py")
    para("About twenty minutes from cold. Then commit and push.")
    para("Two flags are worth knowing:")
    code("python fetch_all_data.py --only air_quality\n"
         "python fetch_all_data.py --export-only")
    para("--only runs a single source. --export-only rebuilds the files the map "
         "reads from data already downloaded, in about a minute, and is the one "
         "you will use most.")

    para("One trap worth knowing.", bold=True)
    para("The map does not read ward_data.json directly. It reads two halves of "
         "it, so that the first paint waits for less. If you rebuild the data "
         "and skip that step, the site keeps serving the previous figures "
         "against a fresh date, and everything will look perfectly reasonable. "
         "The monthly Action does this for you. Doing it by hand:")
    code("python scripts/split_ward_data.py")

    # ── 3 ────────────────────────────────────────────────────────────────────
    h("3. What actually updates when it re-runs", 1)
    para("Re-running does not refresh everything, and it is not meant to. Raw "
         "downloads are kept in a cache that the Action carries between runs, "
         "so a source only downloads again when there is something new to "
         "download. Three different things are going on.")

    h("Refreshes every run", 2)
    para("These are checked against the publisher each time. The NHS registers "
         "send the version they already hold, so an unchanged file costs one "
         "cheap request and no download.")
    table(["Source", "What happens on a re-run"],
          [("GP practices, pharmacies, dental practices",
            "Revalidated against NHS ODS. Downloads only if the register changed."),
           ("Claimant count", "Fetched from NOMIS each run."),
           ("QOF, and the three Fingertips sets", "Fetched from OHID each run."),
           ("Fuel poverty, PTAL", "Fetched each run; both find the current file "
            "from the publisher's own index rather than a saved link."),
           ("Charities and CICs", "Fetched from the Charity Commission each run.")],
          widths=[2.3, 4.0])

    h("Refreshes only when the publisher releases something new", 2)
    para("Cached by period, so the cache is the reason nothing is re-downloaded, "
         "not a fault. A new month or a new year has a new name, so it arrives "
         "on the next run by itself.")
    table(["Source", "What happens on a re-run"],
          [("Street crime", "Cached per month. New months download; old ones do not."),
           ("Air quality", "Cached per pollutant per year. When Defra publish a new "
            "year it downloads, and the run warns you that the year printed on the "
            "map now needs updating."),
           ("Census 2021", "Cached. Fixed until the next census in 2031."),
           ("Boundaries", "Cached. Changes only when ONS revise them.")],
          widths=[2.3, 4.0])

    h("Never refreshes on its own", 2)
    table(["Source", "Why, and what to do"],
          [("Deprivation (IMD 2025)",
            "The processed file is committed and is the source of truth. It is "
            "published about every six years. When a new one lands, drop the raw "
            "file in and re-run that source."),
           ("Green and blue space",
            "Built once from a very large published extract rather than fetched "
            "every month."),
           ("Hospitals",
            "Needs a CSV downloaded by hand. There is no machine-readable source. "
            "Without it the layer is simply empty."),
           ("Your own figures (data/custom)",
            "Only changes when you change the CSV. See section 5.")],
          widths=[2.3, 4.0])

    para("Where to look when something breaks.", bold=True)
    para("data/meta/manifest.json records, for every source, whether the last run "
         "worked, how many rows it wrote, and the error if it failed. The public "
         "methodology page shows the same thing, so a source that has started "
         "failing is visible rather than quietly serving old figures. One source "
         "failing does not stop the others.")

    # ── 4 ────────────────────────────────────────────────────────────────────
    h("4. What will go stale", 1)
    para("None of these are faults. They are dates and identifiers that were "
         "true when they were written.")
    table(["What", "When", "What to do"],
          [("Ward boundaries", "Next boundary review",
            "One line in fetch_all_data.py names the 2025 lookup. Point it at the "
            "new one. Every ward-level figure flows through it."),
           ("Air quality year", "Each autumn",
            "The figures update themselves. The year printed beside them does not. "
            "The run tells you which labels to change."),
           ("Deprivation", "About every 6 years",
            "A new release changes the file name and the columns."),
           ("Census", "2031", "Fixed until then."),
           ("Hospitals", "Any time",
            "Needs a manual CSV, as above.")],
          widths=[1.5, 1.4, 3.4])

    # ── 5 ────────────────────────────────────────────────────────────────────
    h("5. Adding new data", 1)
    para("There are two cases and they are very different amounts of work. Work "
         "out which one you have first.")

    h("You have a spreadsheet of your own figures", 2)
    para("This is the common case, and it needs no programming. Put two files in "
         "the data/custom folder.")
    para("A CSV, where the first column is the area code. Use LSOA21CD for "
         "neighbourhoods or WD25CD for wards:")
    code("LSOA21CD,wellbeing_score\nE01000001,7.2\nE01000002,6.8")
    para("And a JSON file of the same name beside it, describing the columns:")
    code('{\n'
         '  "source": "My team\'s survey, 2026",\n'
         '  "indicators": [{\n'
         '    "column": "wellbeing_score",\n'
         '    "label": "Wellbeing score (1-10)",\n'
         '    "higher_is_worse": false\n'
         '  }]\n'
         '}')
    para("Then run these two, commit, and push:")
    code("python fetch_all_data.py --only custom\npython fetch_all_data.py --export-only")
    para("It appears in the indicator list under Your data, colours the map, and "
         "rolls up from neighbourhoods to wards using the same population "
         "weighting every other source uses. Areas you do not supply are left "
         "blank rather than filled in.")
    para("One field has no default: higher_is_worse.", bold=True)
    para("It decides which end of the colour scale is red. Nothing guesses it, "
         "because a wrong guess does not look like an error. It looks like a "
         "finding. A column without it is refused, with a message saying so.")
    para("There is a worked example in that folder, and a README listing every "
         "field you can set.")

    h("You are adding a national source that refreshes itself", 2)
    para("This is the harder case: a feed somebody else publishes that you want "
         "pulled fresh every month. It touches five places in the code. Copy the "
         "air quality source, which was added exactly this way and is commented "
         "throughout.")
    numbered("Write a function that downloads it and returns a table keyed by "
             "area code. Find the download link from the publisher's index "
             "rather than hardcoding one, because hardcoded links rot.")
    numbered("Add it to the list of sources, which is what --only reads.")
    numbered("Roll it up to wards, beside the other neighbourhood-level sources.")
    numbered("Add it to the neighbourhood file.")
    numbered("Show it in the map: a name, a group, and which end is bad.")

    para("What you do not have to work out.", bold=True)
    para("The colour range, which area sizes have figures, and how much an "
         "indicator actually varies are all measured by the pipeline. Step five "
         "needs a name, a group and a direction; the numbers look after "
         "themselves. If a new indicator does not appear on the map at all, the "
         "missing direction is the first thing to check.")

    # ── 6 ────────────────────────────────────────────────────────────────────
    h("6. Known limitations", 1)
    para("Written down plainly, because a handover that hides them is worse than "
         "no handover.")
    bullet("The browser tests do not all pass. Seven have a failing check, mostly "
           "selectors that went stale as the interface changed. Worth fixing "
           "first: a suite that always fails is one you learn to ignore.")
    bullet("The City of London's 20 wards have no deprivation or health figures. "
           "The wards are tiny and share almost no neighbourhoods, so there is "
           "nothing to average. That is the geography, not a fault.")
    bullet("Five layers cover North West London only: schools, community centres, "
           "libraries, ESOL providers and community interest companies. They are "
           "badged on the map. Blank elsewhere means not collected, not none.")
    bullet("DWP benefit figures are not included. The publicly available ones "
           "stop in 2018. Current figures need a DWP Stat-Xplore account.")
    bullet("Some health figures are published for a whole borough and repeated "
           "for every ward in it. They cannot tell wards apart, and the map says "
           "so beside them.")

    # ── 7 ────────────────────────────────────────────────────────────────────
    h("7. Where everything else is written down", 1)
    table(["File", "What it holds"],
          [("HANDOVER.md", "This document, in the repository, kept up to date with the code."),
           ("fetch_all_data.py", "The pipeline. Long, and written to be read."),
           ("RAW_DATA_SOURCES.md", "Every raw file, its address, and how to fetch it by hand."),
           ("DATA_LICENCES.md", "Licence and required wording for each source. Read before republishing."),
           ("methodology.html", "The public page: every source, its geography and its year."),
           ("data/custom/README.md", "Adding your own figures."),
           ("tests/README.md", "Running the tests.")],
          widths=[2.0, 4.3])

    d.add_paragraph()
    para("Generated from the repository by scripts/build_handover_doc.py. If it "
         "disagrees with the code, the code is right.", italic=True, size=9,
         colour=GREY)

    out.parent.mkdir(parents=True, exist_ok=True)
    d.save(out)
    print(f"wrote {out}  ({out.stat().st_size:,} bytes)")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
